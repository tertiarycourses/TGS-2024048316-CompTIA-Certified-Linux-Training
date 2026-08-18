#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build the WSQ assessment set for 'Application Integration with Docker and Kubernetes' (TGS-2021010366):
  - Written Assessment (SAQ)  — 5 open-ended KNOWLEDGE questions (K1–K5), aligned to the slides
  - Practical Performance (PP) — 4 PRACTICAL tasks (LO1–LO4), aligned to the in-class activities
Each instrument is produced as a Question Paper and a matching Answer Key (4 DOCX total),
all with the WSQ house cover page (same as the Lesson Plan / Learner Guide). Page 1 is the cover;
page 2 carries Trainee Information + Instructions + Grading; the questions/tasks begin on page 3.
Body: Arial 11.
"""
import os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# This script lives in the wsq-assessment skill (.claude/skills/wsq-assessment/) and runs in
# place — it detects the course repo root by walking up to the nearest dir that has a .git
# folder (or both courseware/ and assessment/). Override with env REPO=/path if needed.
def _find_repo():
    env = os.environ.get("REPO")
    if env and os.path.isdir(env):
        return os.path.abspath(env)
    d = os.path.dirname(os.path.abspath(__file__))
    while d != os.path.dirname(d):
        if os.path.isdir(os.path.join(d, ".git")) or \
           (os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "assessment"))):
            return d
        d = os.path.dirname(d)
    return os.getcwd()

REPO = _find_repo()
# prodoc.py (WSQ cover page + version control + page numbers, same as LP/LG) ships with the
# tertiary-lesson-plan skill. Look for it at the project level first, then the user level.
for _cand in (os.path.join(REPO, ".claude/skills/tertiary-lesson-plan"),
              os.path.expanduser("~/.claude/skills/tertiary-lesson-plan")):
    if os.path.exists(os.path.join(_cand, "prodoc.py")):
        sys.path.insert(0, _cand); break
import prodoc  # cover page + version control + page numbers (same as LP/LG)

# ─── EDIT PER COURSE ────────────────────────────────────────────────────────
TITLE       = "CompTIA Certified Linux+ Training"
COURSE_CODE = "TGS-2024048316"
# ────────────────────────────────────────────────────────────────────────────
# The cover page renders prodoc's module-level TGS constant. Override it so the
# assessment cover shows THIS course's ref (works with either prodoc version —
# the older project prodoc has no course_code kwarg).
prodoc.TGS = f"TGS Ref No: {COURSE_CODE}"
OUT   = os.path.join(REPO, "assessment")

# Logos: prefer the course's own courseware/assets, else fall back to the copies bundled
# in this skill (so the assessment builds even outside this project). Replace the course
# logo per course; the Tertiary Infotech logo is the same for every WSQ course.
def _logo(name):
    here = os.path.dirname(os.path.abspath(__file__))
    for p in (os.path.join(REPO, "courseware/assets", name), os.path.join(here, "assets", name)):
        if os.path.exists(p):
            return p
    return None
ORG_LOGO    = _logo("tertiary-infotech-logo.png")
COURSE_LOGO = _logo("comptia-linux-logo.png")   # None if absent → Tertiary-only cover (as LP/LG)

Q_VER, A_VER = "v2", "v2"   # single standardised version across all four files
BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27); GREY = RGBColor(0x55, 0x5B, 0x66)
# Assessments carry the cover page only — no Document Version Control Record.

# ---------------------------------------------------------------- WRITTEN (KNOWLEDGE)
# (criterion, context, question, [model-answer points]) — each traces to the course slides.
WRITTEN = [
 # Each question tests its accredited Knowledge statement (CP_TIPL_comptia_linuxplus v2.1):
 #  K1 Advanced installation and maintenance procedures
 #  K2 Critical components of application support guides
 #  K3 Advanced troubleshooting techniques
 #  K4 Performance analysis of applications
 #  K5 Key factors or considerations in evaluating change requests
 ("K1",
  "Your team routinely installs and maintains applications on Linux servers. Domain 2 (Services and User "
  "Management) covers software management; Domain 1 covers backup before maintenance.",
  "Describe the advanced procedures you would follow to INSTALL and MAINTAIN an application on a Linux server. "
  "Cover: installing with a package manager on both families (apt/dpkg and dnf/rpm), installing a local "
  "package file with its dependencies resolved, installing from source (configure → make → make install), and "
  "the routine maintenance procedures you would schedule — applying updates safely and backing up the "
  "application's data before any change.",
  ["Debian family: apt update && apt install <pkg> (dpkg -i for a local .deb, then apt -f install); RHEL "
   "family: dnf install <pkg> (rpm -i / dnf install ./file.rpm for a local file) — the package manager "
   "resolves dependencies from the repositories.",
   "Local package: sudo apt install ./crm_analytics.deb — apt resolves and installs the dependency chain; "
   "verify with dpkg -l / dpkg -L (or rpm -qa / rpm -ql).",
   "From source: install build tools, then ./configure (checks prerequisites, writes the Makefile), make "
   "(compiles), sudo make install (installs — typically under /usr/local per the FHS).",
   "Maintenance: apply updates on a schedule (apt update/upgrade, dnf upgrade, unattended-upgrades for "
   "security patches); restart and verify the service after (systemctl status / journalctl -u).",
   "Before any change: back up the application data and configs — tar -czf app-$(date +%F).tar.gz, or rsync "
   "-a to a backup target — so the change can be rolled back. "
   "(Slides: Domain 2 — Building/Installing Software & Managing Software Configurations / Lab 11 — Packages / "
   "Lab 6 — Backup & Restore)"]),
 ("K2",
  "A support (training) guide lets any administrator on the team install, operate and troubleshoot an "
  "application consistently — you write one in Lab 24 and for the CRM application in the practical.",
  "You are asked to write the support guide for an application running on a Linux server. Identify the "
  "CRITICAL COMPONENTS the guide must contain and explain why each is needed. Include at least: installation "
  "steps and prerequisites, configuration files and their locations, service operation, log locations, routine "
  "maintenance, and basic troubleshooting with escalation.",
  ["Installation section — prerequisites (packages, resources) and the exact install commands, so a new "
   "administrator can rebuild the application from scratch.",
   "Configuration section — which files control the application and where they live (per the FHS: /etc for "
   "config), with the meaning of the key settings and a copy of the known-good config.",
   "Operation section — how to start/stop/enable the service under systemd (systemctl start/stop/enable/"
   "status) and how to confirm it is healthy (ss -tlnp for its port).",
   "Logging section — where the logs are (/var/log, journalctl -u <unit>) and how to read/filter them, "
   "because logs are the first stop in any incident.",
   "Maintenance section — the update procedure and the backup/restore procedure (tar/rsync) with its schedule.",
   "Troubleshooting section — the common failure symptoms, the checks to run, and when/to whom to escalate. "
   "(Slides: Domain 4 — Lab 24 — documentation & responsible AI / Domain 2 — Lab 12 — systemd / Lab 6 — Backup)"]),
 ("K3",
  "Domain 5 (Troubleshooting) teaches a systematic method plus the per-subsystem techniques and tools.",
  "Describe the systematic troubleshooting method you would apply when an application on a Linux server stops "
  "working. Then name the ADVANCED TECHNIQUES/TOOLS you would use to diagnose each of the following: (a) a "
  "failed systemd service, (b) a storage/disk problem, and (c) a network connectivity problem.",
  ["Method: establish/compare against a baseline → reproduce and scope the problem → gather evidence (logs, "
   "counters) → isolate the layer (service, OS, storage, network) → form and test a hypothesis → remediate → "
   "verify against the baseline and document.",
   "(a) Failed service: systemctl status <unit> (state + last lines), journalctl -u <unit> -p err --since … "
   "(full log), systemd-analyze verify (unit-file errors); check dependencies/ordering with systemctl "
   "list-dependencies.",
   "(b) Storage: df -h / df -i (space AND inode exhaustion), du -sh (what is consuming), lsblk/mount (is it "
   "mounted correctly), dmesg (I/O errors), smartctl -a (disk health), iostat -x (saturation).",
   "(c) Network: ip a / ip route (address & routing), ss -tlnp (is the port listening), ping / traceroute "
   "(reachability & path), dig/nslookup (DNS), tcpdump (what is actually on the wire), plus the firewall "
   "rules (nft list ruleset / firewall-cmd --list-all).",
   "Fix the root cause, not the symptom, and record the finding in the support guide. "
   "(Slides: Domain 5 — Labs 26–28 — Storage / Network / Security Troubleshooting)"]),
 ("K4",
  "Domain 5 covers monitoring vocabulary and the counters that quantify an application's performance.",
  "Explain how you would ANALYSE THE PERFORMANCE of an application on a Linux host. Describe the "
  "performance-analysis loop; for a CPU bottleneck, a memory-pressure problem and a disk-I/O bottleneck name "
  "the command/counter that confirms each; and briefly contrast SLA, SLI and SLO.",
  ["The loop: establish a baseline → observe/reproduce the degradation → identify the bottleneck with the "
   "right counter → remediate → verify recovery against the baseline.",
   "CPU bottleneck: load average vs core count (uptime/nproc), confirmed per-CPU with mpstat -P ALL or top "
   "(%us/%sy); attribute with ps/pidstat.",
   "Memory pressure: free -m shows low available memory and vmstat shows non-zero si/so (swap in/out).",
   "Disk-I/O bottleneck: iostat -x shows high %util and await (processes stuck in D state); pidstat -d "
   "attributes the I/O to a process.",
   "SLI = a measured indicator (e.g. request latency, error rate); SLO = the internal target for an SLI "
   "(e.g. 99.9% of requests < 200 ms); SLA = the external, contractual promise (often with penalties) built "
   "on SLOs. (Slides: Domain 5 — Key Concepts / Lab 25 — Monitoring / Lab 29 — Performance)"]),
 ("K5",
  "Applications evolve through change requests, which must be evaluated before they are implemented — as you "
  "do for the CRM application in Lab 21 and in the practical assessment.",
  "You receive three change requests for a CRM application on your Linux server: real-time data "
  "synchronisation, a new reporting module, and a database upgrade. What KEY FACTORS would you consider in "
  "evaluating whether each request is VALID and FEASIBLE? Include the commands you would use to check the "
  "host's capacity, and the risk controls you would insist on before approving a change.",
  ["Validity: does the request serve a real business/user need (check against user feedback), and is it in "
   "scope for the application?",
   "Resource feasibility: check the host's capacity against the request's requirements — nproc (CPU cores), "
   "free -m (available memory), df -h / df -m (free disk) — as scripted in evaluate_changes.sh (Lab 21).",
   "Compatibility & dependencies: OS/package versions, database compatibility for the upgrade, impact on the "
   "other components.",
   "Security impact: new open ports/services, new data flows, required firewall/hardening changes.",
   "Risk controls before approval: full backup with a tested rollback plan (tar/rsync), test in a staging "
   "environment first, an agreed maintenance/downtime window, and monitoring after the change to verify "
   "success.",
   "Outcome: rank each request feasible / not feasible on the current host, and propose the enhancement path "
   "to the developers. (Slides: Domain 4 — Lab 21 — Bash Scripting (evaluate_changes.sh) / Lab 6 — Backup / "
   "Domain 5 — Monitoring)"]),
]

# ---------------------------------------------------------------- PRACTICAL (ACTIVITY-BASED)
SCENARIO = (
 "You have been hired as a junior Linux administrator at CompTech Solutions, a medium-sized tech company that "
 "has recently migrated its infrastructure to Linux to improve security and scalability. Your supervisor has "
 "tasked you with keeping the company's CRM application running smoothly: evaluating change requests, producing "
 "a short training guide, troubleshooting application issues from logs, analysing system performance, and "
 "deploying a new feature package. Complete the four tasks below; each mirrors a hands-on lab you did in class. "
 "For each task, paste your commands, your script and output snapshots, and any text-file contents as evidence.")

BOX_CAP = "Paste your commands, script and output snapshots, and text-file contents in the box below"
PRACTICAL = [
 ("Task 1", "LO1 · A7, A2",
  "Evaluate change requests and write a training guide. You've been given several change requests for the CRM "
  "application — real-time data synchronisation, a new reporting module, and a database upgrade. "
  "Part A — Write a Bash script (evaluate_changes.sh) that inspects the host's resources (CPU cores, available "
  "memory and free disk space) and, using conditional logic and a loop, prints which requests are feasible on "
  "the current system. Use a shebang, set -euo pipefail, variables, an if/[[ ]] test and a for loop, then "
  "confirm the script is clean with shellcheck. "
  "Part B — Write a short training guide (crm_training.txt) covering how to install and maintain the CRM "
  "application. Provide your script, the output of running it, and the first lines of the training guide. "
  "(Aligns to Lab 21 — Bash Scripting; Lab 24 — Responsible AI / documentation.)",
  BOX_CAP,
  "Part A — evaluate_changes.sh:\n"
  "#!/usr/bin/env bash\n"
  "set -euo pipefail\n"
  "cores=$(nproc)\n"
  "mem_mb=$(free -m | awk '/^Mem:/{print $7}')            # available RAM (MB)\n"
  "disk_mb=$(df -m --output=avail / | tail -1 | tr -d ' ')\n"
  "# name:cores:mem_mb:disk_mb required\n"
  "requests=(\"realtime-sync:2:1024:5000\" \"reporting:1:512:2000\" \"db-upgrade:4:2048:10000\")\n"
  "for r in \"${requests[@]}\"; do\n"
  "  IFS=':' read -r name c m d <<<\"$r\"\n"
  "  if [[ $cores -ge $c && $mem_mb -ge $m && $disk_mb -ge $d ]]; then\n"
  "    echo \"FEASIBLE: $name\"\n"
  "  else\n"
  "    echo \"NOT FEASIBLE: $name (needs ${c} cores, ${m}MB RAM, ${d}MB disk)\"\n"
  "  fi\n"
  "done\n"
  "# Run & lint:\n"
  "chmod +x evaluate_changes.sh && ./evaluate_changes.sh\n"
  "shellcheck evaluate_changes.sh          # clean report\n\n"
  "Part B — crm_training.txt (first lines):\n"
  "CRM Application - Install & Maintenance Guide\n"
  "1. Install:  sudo apt install ./crm_analytics.deb\n"
  "2. Start:    sudo systemctl enable --now crm\n"
  "3. Logs:     journalctl -u crm -f\n\n"
  "Award the mark where the candidate uses resource checks (nproc / free / df), a conditional plus a loop, a "
  "shellcheck-clean script, and a clear training guide. (Lab 21 — Bash Scripting; Lab 24 — documentation.)"),
 ("Task 2", "LO2 · A3, A4",
  "Troubleshoot the CRM from its logs. The CRM application has been experiencing slowdowns. Analyse the provided "
  "log file crm_app.log using Linux text-processing tools to find the errors and patterns behind the slowdown. "
  "Filter the log for errors, count how often each error type occurs, and identify the busiest time window. "
  "Document your findings and the steps you would take to resolve them in findings.txt. Provide the commands you "
  "used, the output showing the relevant errors, and your findings file. "
  "(Aligns to Lab 5 — Shell Operations & Text Processing; Lab 25 — Monitoring / log aggregation.)",
  BOX_CAP,
  "Filter and quantify the errors:\n"
  "grep -iE \"error|fail|timeout\" crm_app.log | head\n"
  "grep -c -i error crm_app.log                      # total error count\n"
  "awk '/ERROR/{print $NF}' crm_app.log | sort | uniq -c | sort -rn   # top error types\n"
  "# Busiest window (group by hour):\n"
  "awk '{print $2}' crm_app.log | cut -d: -f1 | sort | uniq -c | sort -rn | head\n"
  "# For a running service, the journald equivalent:\n"
  "journalctl -u crm --since \"1 hour ago\" -p err --no-pager\n\n"
  "findings.txt:\n"
  "- Most frequent error: database connection timeout (NN occurrences).\n"
  "- Peak error volume between 14:00-15:00.\n"
  "- Proposed fix: increase the DB connection-pool size, add the missing index, then restart crm.\n\n"
  "Award the mark for correct grep / awk / sort / uniq filtering, identifying the dominant error and the peak "
  "window, and a clear findings file with remediation steps. (Lab 5 — Text Processing; Lab 25 — Monitoring.)"),
 ("Task 3", "LO3 · A5, A6",
  "Analyse performance and optimise. Review the provided performance_report.txt for the CRM host and identify at "
  "least three underlying issues contributing to poor performance (consider CPU, memory, disk I/O and "
  "processes). Then read user_feedback.txt and propose concrete changes to optimise the application. List the "
  "commands you would run on a live host to confirm each issue, write the three issues and your analysis to "
  "perf_issues.txt, and record your proposed changes. "
  "(Aligns to Lab 29 — Performance Troubleshooting; Lab 10 — Processes, Jobs & Scheduling.)",
  BOX_CAP,
  "Confirm each issue on a live host:\n"
  "uptime; mpstat -P ALL 1 3               # CPU: load > cores, a core pegged at 100%\n"
  "free -m; vmstat 1 5                      # Memory: low available; si/so > 0 = swapping\n"
  "iostat -x 1 3; pidstat -d 1 3            # Disk I/O: high %util/await; top writer PID\n"
  "ps -eo pid,pcpu,pmem,stat,cmd --sort=-pcpu | head   # top consumers; D = I/O wait\n\n"
  "perf_issues.txt (at least three):\n"
  "1. CPU saturation - load average exceeds the core count (mpstat shows a thread at 100%).\n"
  "2. Memory pressure - free shows little available RAM; vmstat si/so are non-zero (swapping).\n"
  "3. Disk I/O bottleneck - iostat %util near 100% with high await; a process stuck in D state.\n\n"
  "Proposed changes (from user_feedback.txt):\n"
  "- Add an index/cache to cut database CPU; fix the memory leak or add RAM; move the DB to faster storage; "
  "renice the heavy batch job so it stops starving the app.\n\n"
  "Award the mark for mapping each symptom to the right counter (load vs mpstat, free vs vmstat si/so, iostat "
  "%util/await), three genuine issues, and feedback-driven optimisations. (Lab 29 — Performance; Lab 10 — Processes.)"),
 ("Task 4", "LO4 · A1, A8",
  "Deploy and test a new feature package. A new data-analytics feature is shipped as a Debian package, "
  "crm_analytics.deb. Install it, configure it to run as a service, and test it. Show the commands to install "
  "the package (resolving dependencies), enable and start it under systemd, confirm it is active and listening, "
  "and read its logs. After testing, write a short performance assessment plus one suggested enhancement to "
  "analytics_assessment.txt. Provide the install commands and output, the configure/test commands, and your "
  "assessment file. (Aligns to Lab 11 — Software & Package Management; Lab 12 — Service Management with systemd.)",
  BOX_CAP,
  "Install the .deb (apt resolves dependencies):\n"
  "sudo apt install ./crm_analytics.deb     # or: sudo dpkg -i crm_analytics.deb; sudo apt -f install\n"
  "dpkg -l | grep crm-analytics             # confirm installed\n"
  "dpkg -L crm-analytics | head             # files the package placed\n\n"
  "Configure and run it as a service:\n"
  "sudo systemctl daemon-reload\n"
  "sudo systemctl enable --now crm-analytics\n"
  "systemctl status crm-analytics --no-pager   # active (running)\n"
  "ss -tlnp | grep crm-analytics               # confirm it is listening\n"
  "journalctl -u crm-analytics -n 20 --no-pager\n\n"
  "analytics_assessment.txt:\n"
  "- Feature responds within acceptable latency; CPU/memory footprint reasonable under test load.\n"
  "- Suggested enhancement: add a systemd timer for scheduled report generation, or cache results.\n\n"
  "Award the mark for a correct .deb install with dependency resolution (apt install ./file.deb, or dpkg -i + "
  "apt -f install), enabling/starting under systemd, verifying active + listening, reading journald logs, and a "
  "sensible assessment. (Lab 11 — Packages; Lab 12 — systemd.)"),
]

# ---------------------------------------------------------------- doc helpers
def base_doc():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Arial"; n.font.size = Pt(11)
    return doc

def para(doc, text, size=11, bold=False, italic=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if align is not None: p.alignment = align
    return p

def heading(doc, text, size=13):
    para(doc, text, size=size, bold=True, color=BRAND, after=6, before=8)

def answer_box(doc, lines=None, code=None, height_pt=90):
    """1x1 bordered box. `lines` → bullet-style model answer; `code` → monospace
    code/YAML/command block (indentation preserved); neither → empty answer space."""
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.paragraphs[0].text = ""
    if code:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in code.split("\n"):
            b = cell.add_paragraph(style=None)
            b.paragraph_format.space_after = Pt(0); b.paragraph_format.space_before = Pt(0)
            rr = b.add_run(ln if ln else " ")
            rr.font.name = "Consolas"; rr.font.size = Pt(9)
            rr._element.rPr.rFonts.set(qn('w:cs'), "Consolas")
            wt = rr._element.find(qn('w:t'))
            if wt is not None: wt.set(qn('xml:space'), 'preserve')
    elif lines:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in lines:
            b = cell.add_paragraph(style=None); b.paragraph_format.left_indent = Inches(0.15)
            rr = b.add_run("•  " + ln); rr.font.size = Pt(10.5)
    else:
        # empty answer space
        tr = t.rows[0]._tr
        trPr = tr.get_or_add_trPr(); trh = OxmlElement('w:trHeight')
        trh.set(qn('w:val'), str(int(height_pt*20))); trh.set(qn('w:hRule'), 'atLeast'); trPr.append(trh)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

FILL_GAP = 6    # extra space below each fill-in line (paired with double line spacing for writing room)

def candidate_block(doc):
    heading(doc, "Trainee Information")
    for label in ["Trainee Name (as per NRIC): ______________________________________",
                  "Last 3 digits and alphabet of NRIC/FIN: ____________________",
                  "Date: ____________________"]:
        p = para(doc, label, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

# Assessment briefing (from the course slides — "Briefing for Assessment").
BRIEFING = [
    "Place phones and other materials under the table or on the floor.",
    "No photos or recording of assessment scripts.",
    "No discussion during the assessment.",
    "Use a black/blue pen for hard-copy assessments.",
    "No liquid paper / correction tape.",
    "Scripts are collected when time is up.",
]

LMS_URL = "https://lms-tms.tertiaryinfotech.com/"

def add_hyperlink(p, url, text):
    """Add a real clickable Word hyperlink (blue, underlined) to paragraph p."""
    r_id = p.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22"); rPr.append(sz)  # 11pt
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run); p._p.append(link)
    return link

def instructions(doc, minutes_text):
    heading(doc, "Instructions to Candidate")
    # None marks the upload instruction, which carries a clickable LMS hyperlink.
    items = [
        "This is an individual exercise.",
        "This is an open-book assessment.",
        f"A total of {minutes_text} is given to complete this assessment.",
        None,
    ] + BRIEFING
    for i, s in enumerate(items, 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        if s is None:
            p.add_run(f"{i}.  Complete your answers on the document provided and "
                      "upload the completed answers to the LMS at ").font.size = Pt(11)
            add_hyperlink(p, LMS_URL, LMS_URL)
            p.add_run(".").font.size = Pt(11)
        else:
            p.add_run(f"{i}.  {s}").font.size = Pt(11)

def grading(doc, what):
    heading(doc, "Grading")
    para(doc, what, size=11, after=12)
    for ln in ["Grade: _______  (C / NYC)",
               "Assessor Name: __________________________   Assessor NRIC: ________________",
               "Date: ________________________                    Signature: ____________________"]:
        p = para(doc, ln, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

def finish(doc, path):
    prodoc.add_page_numbers(doc); prodoc.enable_update_fields(doc)
    doc.save(path); print("  saved:", os.path.basename(path))

# ---------------------------------------------------------------- builders
def build_wa(answers):
    doc = base_doc()
    kind = "Written Assessment (SAQ) — Answer Key" if answers else "Written Assessment (SAQ)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Written Assessment (SAQ)" if answers else "Written Assessment (SAQ)",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        # Page 2 — candidate information, instructions and grading; questions begin on the next page.
        candidate_block(doc); instructions(doc, "1 hour")
        grading(doc, "Candidate has answered all written questions and demonstrated the underpinning "
                     "knowledge required for the course learning outcomes.")
        page_break(doc)
    para(doc, "Short-Answer Questions (Knowledge)", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Answer all questions in your own words. Each question tests underpinning knowledge covered in the "
              "course slides.", size=10.5, italic=True, color=GREY, after=8)
    for i, (crit, ctx, q, pts) in enumerate(WRITTEN, 1):
        para(doc, f"Question {i}:", size=11.5, bold=True, after=2, before=6)
        para(doc, ctx, size=11, after=3)
        para(doc, f"{q}  ({crit})", size=11, bold=True, after=4)
        answer_box(doc, lines=pts if answers else None)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to WA (SAQ) - {TITLE} - {suffix}.docx" if answers
            else f"WA (SAQ) - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

def build_pp(answers):
    doc = base_doc()
    kind = "Practical Performance (PP) — Answer Key" if answers else "Practical Performance (PP)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Practical Performance Assessment" if answers else "Practical Performance Assessment",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        # Page 2 — candidate information, instructions and grading; the problem begins on the next page.
        candidate_block(doc); instructions(doc, "1 hour")
        grading(doc, "Candidate has successfully completed all PP tasks and can explain the overall "
                     "functions and features used to achieve them.")
        page_break(doc)
    para(doc, "Practical Problem", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Scenario", size=11.5, bold=True, after=2)
    para(doc, SCENARIO, size=11, after=8)
    for label, crit, prompt, cap, pts in PRACTICAL:
        para(doc, f"{label} ({crit}):", size=11.5, bold=True, after=2, before=6)
        para(doc, prompt, size=11, after=3)
        para(doc, cap, size=10.5, italic=True, color=GREY, after=4)
        answer_box(doc, code=pts if answers else None, height_pt=150)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to PP Assessment - {TITLE} - {suffix}.docx" if answers
            else f"PP Assessment - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

if __name__ == "__main__":
    print("Building WSQ assessment set…")
    build_wa(answers=False); build_wa(answers=True)
    build_pp(answers=False); build_pp(answers=True)
    print(f"Done. WA: {len(WRITTEN)} questions · PP: {len(PRACTICAL)} tasks.")
